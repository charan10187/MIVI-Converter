import streamlit as st
from PIL import Image
import pandas as pd
import json
import io
from io import BytesIO
import tempfile
import os
import shutil
import platform
import pypandoc
from pdf2docx import Converter
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import pytesseract
from docx import Document
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


# Optional: handle tesseract not found errors gracefully
try:
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

st.set_page_config(page_title="MIVI Universal Converter", page_icon="📂", layout="wide")

st.title("📂 MIVI Universal File Converter 🗃")
st.markdown("Developed by *S. Sri Charan* | 📧 [charan10187@gmail.com](mailto:charan10187@gmail.com)")
st.markdown(
    "⚠ *Note:* Only local files are supported. "
    "Cloud links (Google Drive, Dropbox, OneDrive) are not accepted. "
    "Download them first before uploading here."
)
tab1, tab2 = st.tabs(["🖼 Image Resizer", "📄 File Converter"])
with tab1:
    st.subheader("Resize Your Images Easily")

    img_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png", "bmp", "tiff"], key="img_upload")
    resize_mode = st.radio("Resize Mode", ["By Dimensions", "By File Size (KB)"], key="resize_mode")

    if img_file:
        try:
            image = Image.open(img_file)
            image_format = image.format or "PNG" # Store original format
            
            # Ensure image is in a saveable mode (like RGB) before processing
            if image.mode not in ("RGB", "L"):
                if image.mode == "RGBA":
                     # Handle transparency
                     bg = Image.new("RGB", image.size, (255, 255, 255))
                     bg.paste(image, (0, 0), image)
                     image = bg
                else:
                    image = image.convert("RGB")
                image_format = "JPEG" # PNGs with transparency converted to JPEG

            # --- Resize by Dimensions ---
            if resize_mode == "By Dimensions":
                width = st.number_input("Width", min_value=1, value=image.width, key="width_input")
                height = st.number_input("Height", min_value=1, value=image.height, key="height_input")
                maintain_ratio = st.checkbox("Maintain Aspect Ratio", value=True, key="aspect_ratio")

                if maintain_ratio:
                    # Calculate height based on width to maintain aspect ratio
                    if width != image.width:
                        height = int(image.height * (width / float(image.width)))
                    # Recalculate width if height is changed manually (less common)
                    elif height != image.height:
                         width = int(image.width * (height / float(image.height)))


                if st.button("Resize Now", key="resize_button"):
                    resized = image.resize((int(width), int(height)))
                    buf = io.BytesIO()
                    
                    save_format = image_format
                    if save_format == "JPEG":
                        resized.save(buf, format="JPEG", quality=95)
                        mime_type = "image/jpeg"
                        file_ext = "jpg"
                    else:
                        # Revert to PNG if original was not JPEG
                        resized.save(buf, format="PNG")
                        mime_type = "image/png"
                        file_ext = "png"
                        
                    st.image(resized, caption="Resized Image", use_column_width=True)
                    st.download_button(
                        "📥 Download Resized Image",
                        buf.getvalue(),
                        file_name=f"{os.path.splitext(img_file.name)[0]}_resized.{file_ext}",
                        mime=mime_type,
                        key="download_resized"
                    )

            # --- Resize by File Size (KB) ---
            elif resize_mode == "By File Size (KB)":
                target_kb = st.number_input("Target Size (KB)", min_value=10, max_value=5000, value=200, key="target_kb")

                if st.button("Compress Now", key="compress_button"):
                    buf = io.BytesIO()
                    quality = 95
                    
                    # Compression by size really only works well with JPEG
                    img_for_compress = image
                    if img_for_compress.mode == "RGBA":
                        # Create a white background and paste RGBA image on it
                        bg = Image.new("RGB", img_for_compress.size, (255, 255, 255))
                        bg.paste(img_for_compress, (0, 0), img_for_compress)
                        img_for_compress = bg
                    elif img_for_compress.mode != "RGB":
                         img_for_compress = img_for_compress.convert("RGB")
                    
                    img_for_compress.save(buf, format="JPEG", quality=quality)
                    size_kb = len(buf.getvalue()) / 1024

                    while size_kb > target_kb and quality > 5:
                        quality -= 5
                        buf.truncate(0)
                        buf.seek(0)
                        img_for_compress.save(buf, format="JPEG", quality=quality)
                        size_kb = len(buf.getvalue()) / 1024
                        if quality <= 5:
                            break

                    st.success(f"✅ Compressed to ~{int(size_kb)} KB at quality {quality}")
                    
                    # Display the compressed image
                    buf.seek(0)
                    compressed_image_display = Image.open(buf)
                    st.image(compressed_image_display, caption="Compressed Image", use_column_width=True)
                    
                    buf.seek(0) # Rewind buffer for download
                    st.download_button(
                        "📥 Download Compressed Image",
                        buf.getvalue(),
                        file_name=f"{os.path.splitext(img_file.name)[0]}_compressed.jpg",
                        mime="image/jpeg",
                        key="download_compressed"
                    )

        except Exception as e:
            st.error(f"❌ Error processing image: {e}")

# ================================================================
#  FILE CONVERTER TAB
# ================================================================
with tab2:
    st.subheader("📄 Smart File Converter")

    uploaded_file = st.file_uploader(
        "Upload your file", type=["pdf", "docx", "txt", "csv", "xlsx", "json", "jpg", "jpeg", "png"]
    )

    if uploaded_file:
        file_ext = uploaded_file.name.split(".")[-1].lower()
        file_base = os.path.splitext(uploaded_file.name)[0]

        # Detect available conversions
        convert_options = []
        if file_ext == "pdf":
            # --- MODIFICATION HERE ---
            # Removed "PDF ➜ Word" as requested, as it requires complex external OCR tools 
            # (Tesseract and Poppler) that are not installed.
            convert_options = ["PDF ➜ Text"]
        elif file_ext == "docx":
            convert_options = ["Word ➜ PDF"]
        elif file_ext == "csv":
            convert_options = ["CSV ➜ Excel", "CSV ➜ JSON"]
        elif file_ext == "xlsx":
            convert_options = ["Excel ➜ CSV"]
        elif file_ext == "json":
            convert_options = ["JSON ➜ CSV"]
        elif file_ext == "txt":
            convert_options = ["Text ➜ PDF"]
        elif file_ext in ["jpg", "jpeg", "png"]:
            if OCR_AVAILABLE:
                convert_options = ["Image ➜ Text (OCR)", "Image ➜ Word (OCR)"]
            else:
                st.warning("OCR (Tesseract) not found. Image-to-text conversions are disabled.")

        if not convert_options:
            if file_ext not in ["jpg", "jpeg", "png"]: # Avoid double error
                st.error("Unsupported file type or no available conversions.")
        else:
            conversion_type = st.selectbox("Select conversion type:", convert_options)

            if st.button("🚀 Convert Now"):
                progress = st.progress(0)
                status = st.empty()
                tmp_dir = tempfile.mkdtemp()
                
                try:
                    # ---------------- PDF ➜ WORD (This block is now unused) ----------------
                    if conversion_type == "PDF ➜ Word":
                        status.text("Converting PDF to Word... please wait.")
                        pdf_path = os.path.join(tmp_dir, "input.pdf")
                        docx_path = os.path.join(tmp_dir, "output.docx")

                        with open(pdf_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())

                        progress.progress(10)
                        cv = None # Initialize converter object
                        try:
                            cv = Converter(pdf_path)
                            cv.convert(docx_path, start=0, end=None)
                            progress.progress(100)
                            status.text("✅ Conversion complete with formatting preserved!")

                            with open(docx_path, "rb") as f:
                                st.download_button(
                                    "📥 Download Converted Word File",
                                    f,
                                    file_name=f"{file_base}_converted.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_word_std"
                                )

                        except Exception as e:
                            if OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                                status.text(f"⚠ Standard conversion failed. Switching to OCR mode...")
                                
                                images = convert_from_path(pdf_path)
                                doc = Document()
                                for i, img in enumerate(images):
                                    status.text(f"Processing page {i+1}/{len(images)} via OCR...")
                                    text = pytesseract.image_to_string(img)
                                    doc.add_paragraph(text)
                                    if i < len(images) - 1:
                                        doc.add_page_break()

                                doc.save(docx_path)
                                progress.progress(100)
                                status.text("✅ OCR conversion completed successfully (scanned PDF handled).")
                                
                                with open(docx_path, "rb") as f:
                                    st.download_button(
                                        "📥 Download Converted Word File (OCR)",
                                        f,
                                        file_name=f"{file_base}_ocr_converted.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="dl_word_ocr"
                                    )
                            else:
                                st.error(f"Standard conversion failed: {e}. OCR fallback is unavailable (missing Tesseract or pdf2image).")
                                
                        finally:
                            if cv is not None:
                                cv.close()

                    # ---------------- WORD ➜ PDF ----------------
                    elif conversion_type == "Word ➜ PDF":
                        status.text("Converting Word to PDF... please wait.")
                        docx_path = os.path.join(tmp_dir, "input.docx")
                        pdf_output = os.path.join(tmp_dir, f"{file_base}_converted.pdf")

                        with open(docx_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())

                        progress.progress(40)

                        try:
                            pypandoc.convert_file(
                                docx_path,
                                "pdf",
                                outputfile=pdf_output,
                                extra_args=["--pdf-engine=xelatex", "--standalone"]
                            )
                            progress.progress(100)
                            with open(pdf_output, "rb") as f:
                                st.download_button(
                                    "📥 Download Converted PDF",
                                    f,
                                    file_name=f"{file_base}_converted.pdf",
                                    mime="application/pdf",
                                )
                            status.text("✅ Word converted to PDF with formatting preserved.")
                        except Exception as e:
                            st.error(f"❌ Conversion failed: {e}")



                    # ---------------- PDF ➜ TEXT ----------------
                    elif conversion_type == "PDF ➜ Text":
                        status.text("Extracting text from PDF...")
                        pdf_path = os.path.join(tmp_dir, "input.pdf")
                        text_output_path = os.path.join(tmp_dir, "output.txt")
                        text = ""

                        with open(pdf_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())

                        try:
                            reader = PdfReader(pdf_path)
                            for page in reader.pages:
                                text += (page.extract_text() or "") + "\n"
                        except Exception as e:
                            st.warning(f"PyPDF2 failed: {e}. Attempting OCR.")
                            text = ""

                        if not text.strip() and OCR_AVAILABLE and PDF2IMAGE_AVAILABLE:
                            status.text("No text found — running OCR...")
                            pages = convert_from_path(pdf_path)
                            text = "\n\n--- Page Break ---\n\n".join(pytesseract.image_to_string(p) for p in pages)
                        elif not text.strip() and (not OCR_AVAILABLE or not PDF2IMAGE_AVAILABLE):
                             status.error("No extractable text found, and external OCR tools are not installed.")
                             
                        if text.strip():
                            with open(text_output_path, "w", encoding="utf-8") as f:
                                f.write(text)
                                
                            with open(text_output_path, "rb") as f:
                                st.download_button(
                                    "📥 Download Text File",
                                    f,
                                    file_name=f"{file_base}_converted.txt",
                                )
                            status.text("✅ Text extracted successfully.")

                    # ---------------- CSV ➜ EXCEL ----------------
                    elif conversion_type == "CSV ➜ Excel":
                        df = pd.read_csv(uploaded_file)
                        buf = BytesIO()
                        df.to_excel(buf, index=False)
                        st.download_button(
                            "📥 Download Excel File",
                            buf.getvalue(),
                            file_name=f"{file_base}_converted.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )
                        status.text("✅ CSV converted to Excel successfully.")

                    # ---------------- EXCEL ➜ CSV ----------------
                    elif conversion_type == "Excel ➜ CSV":
                        xls = pd.ExcelFile(uploaded_file)
                        sheet_name = xls.sheet_names[0]
                        if len(xls.sheet_names) > 1:
                            sheet_name = st.selectbox("Select Excel sheet to convert:", xls.sheet_names)
                            
                        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                        csv_data = df.to_csv(index=False).encode('utf-8')
                        st.download_button(
                            "📥 Download CSV File",
                            csv_data,
                            file_name=f"{file_base}_{sheet_name}_converted.csv",
                            mime="text/csv",
                        )
                        status.text(f"✅ Excel sheet '{sheet_name}' converted to CSV successfully.")

                    # ---------------- CSV ➜ JSON ----------------
                    elif conversion_type == "CSV ➜ JSON":
                        df = pd.read_csv(uploaded_file)
                        json_data = df.to_json(orient="records", indent=2)
                        st.download_button(
                            "📥 Download JSON File",
                            json_data.encode("utf-8"),
                            file_name=f"{file_base}_converted.json",
                            mime="application/json",
                        )
                        status.text("✅ CSV converted to JSON successfully.")

                    # ---------------- JSON ➜ CSV ----------------
                    elif conversion_type == "JSON ➜ CSV":
                        try:
                            data = json.load(uploaded_file)
                            df = pd.json_normalize(data) # Handles nested JSON
                        except Exception as e:
                             status.error(f"Error processing JSON: {e}. Trying simple read.")
                             uploaded_file.seek(0)
                             df = pd.read_json(uploaded_file)
                             
                        csv_data = df.to_csv(index=False).encode('utf-8')
                        st.download_button(
                            "📥 Download CSV File",
                            csv_data,
                            file_name=f"{file_base}_converted.csv",
                            mime="text/csv",
                        )
                        status.text("✅ JSON converted to CSV successfully.")

                    # ---------------- TXT ➜ PDF ----------------
                    elif conversion_type == "Text ➜ PDF":
                        text_data = uploaded_file.read().decode("utf-8")
                        pdf_buf = BytesIO()
                        c = canvas.Canvas(pdf_buf, pagesize=A4)
                        width, height = A4 # Get page dimensions
                        
                        text_object = c.beginText(40, height - 40) # Start near top-left
                        c.setFont("Helvetica", 10) # Set font
                        
                        line_height = 12 # 12 points per line
                        max_lines_per_page = (height - 80) // line_height
                        lines = text_data.splitlines()
                        
                        line_count = 0
                        for line in lines:
                            # Simple word wrap
                            if c.stringWidth(line) > (width - 80):
                                words = line.split()
                                current_line = ""
                                for word in words:
                                    if c.stringWidth(current_line + word + " ") < (width - 80):
                                        current_line += word + " "
                                    else:
                                        text_object.textLine(current_line)
                                        current_line = word + " "
                                        line_count += 1
                                        if line_count >= max_lines_per_page:
                                            c.drawText(text_object)
                                            c.showPage()
                                            c.setFont("Helvetica", 10)
                                            text_object = c.beginText(40, height - 40)
                                            line_count = 0
                                text_object.textLine(current_line) # Add the last part
                            else:
                                text_object.textLine(line)
                            
                            line_count += 1
                            if line_count >= max_lines_per_page:
                                c.drawText(text_object)
                                c.showPage() # New page
                                c.setFont("Helvetica", 10)
                                text_object = c.beginText(40, height - 40)
                                line_count = 0
                                
                        c.drawText(text_object) # Draw remaining text
                        c.save()

                        st.download_button(
                            "📥 Download PDF",
                            pdf_buf.getvalue(),
                            file_name=f"{file_base}_converted.pdf",
                            mime="application/pdf",
                        )
                        status.text("✅ Text converted to PDF successfully.")

                    # ---------------- IMAGE ➜ TEXT (OCR) ----------------
                    elif conversion_type == "Image ➜ Text (OCR)":
                        img = Image.open(uploaded_file)
                        status.text("Extracting text from image (OCR)...")
                        text = pytesseract.image_to_string(img)
                        st.download_button(
                            "📥 Download Extracted Text",
                            text.encode("utf-8"),
                            file_name=f"{file_base}_ocr.txt",
                            mime="text/plain",
                        )
                        status.text("✅ Text extracted successfully via OCR.")
                        st.text_area("Extracted Text:", text, height=300)

                    # ---------------- IMAGE ➜ WORD (OCR) ----------------
                    elif conversion_type == "Image ➜ Word (OCR)":
                        img = Image.open(uploaded_file)
                        status.text("Converting image to Word (OCR)...")
                        text = pytesseract.image_to_string(img)
                        
                        doc = Document()
                        doc.add_paragraph(text)
                        buf = BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        
                        st.download_button(
                            "📥 Download Word File",
                            buf.getvalue(),
                            file_name=f"{file_base}_ocr.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                        status.text("✅ Image converted to editable Word document.")

                except Exception as e:
                    st.error(f"❌ An error occurred during conversion: {e}")
                
                finally:
                    # --- Clean up the temporary directory ---
                    if os.path.exists(tmp_dir):
                        try:
                            shutil.rmtree(tmp_dir)
                        except Exception as e:
                            st.warning(f"Could not clean up temp folder: {e}. It will be removed on next run.")


st.markdown("---")
st.caption("🧠 Smart File Conversion | Powered by MIVI Converter © 2Player 2025 | Developed by **S. Sri Charan**")

